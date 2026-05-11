# ─── Imports ──────────────────────────────────────────────────────────────────
from fastapi import FastAPI, UploadFile, File, BackgroundTasks, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import shutil
import uuid

from pypdf import PdfReader
from PIL import Image
import pytesseract
import docx

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS


import requests
import os
from dotenv import load_dotenv
from groq import Groq
load_dotenv()
groq_api_key = os.environ.get("GROQ_API_KEY")
client = Groq(api_key=groq_api_key)
from langchain_huggingface import HuggingFaceEmbeddings
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
# ─── App setup ────────────────────────────────────────────────────────────────
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Global state ─────────────────────────────────────────────────────────────
vector_db_store: dict = {}  # { session_id: FAISS vector_db }

# { file_id: { "status": "processing"|"ready"|"error", "message": str, "chunks": int } }
upload_status: dict = {}


# ─── Request models ───────────────────────────────────────────────────────────
class HistoryMessage(BaseModel):
    role: str        # "user" | "assistant"
    content: str

class AskRequest(BaseModel):
    query: str
    history: Optional[List[HistoryMessage]] = []
    tutor_mode: Optional[bool] = False
    system_prompt: Optional[str] = None
    session_id: Optional[str] = None  # ← per-user isolation


# ─── Helpers ──────────────────────────────────────────────────────────────────
def extract_text(file_location: str, filename: str) -> str:
    text = ""
    if filename.endswith(".pdf"):
        reader = PdfReader(file_location)
        for page in reader.pages:
            text += page.extract_text() or ""
    elif filename.endswith((".jpg", ".jpeg", ".png")):
        image = Image.open(file_location)
        text = pytesseract.image_to_string(image)
    elif filename.endswith(".docx"):
        doc = docx.Document(file_location)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text


def extract_text_with_pages(file_location: str, filename: str) -> List[dict]:
    pages = []
    if filename.endswith(".pdf"):
        reader = PdfReader(file_location)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append({"page": i + 1, "text": page_text})
    elif filename.endswith((".jpg", ".jpeg", ".png")):
        image = Image.open(file_location)
        text = pytesseract.image_to_string(image)
        pages.append({"page": 1, "text": text})
    elif filename.endswith(".docx"):
        doc = docx.Document(file_location)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        pages.append({"page": 1, "text": full_text})
    return pages

def format_history(history: List[HistoryMessage], max_turns: int = 5) -> str:
    if not history:
        return ""
    recent = history[-max_turns:]
    lines = []
    for msg in recent:
        label = "User" if msg.role == "user" else "Assistant"
        lines.append(f"{label}: {msg.content}")
    return "\n".join(lines)


def score_confidence(context: str, answer: str) -> str:
    if len(context.strip()) < 50:
        return "low"
    uncertainty_signals = [
        "i'm not sure", "i don't know", "cannot find",
        "not found", "unclear", "i am not certain",
    ]
    if any(sig in answer.lower() for sig in uncertainty_signals):
        return "medium"
    return "high"


def call_ollama(prompt: str) -> str:
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2048,
        temperature=0.4,
    )
    return response.choices[0].message.content.strip()


def detect_task(query: str) -> str:
    q = query.lower()
    if any(w in q for w in ["study note", "study notes", "notes", "create note", "make note", "concise note"]):
        return "notes"
    if any(w in q for w in ["summarize", "summary", "summarise", "overview", "brief"]):
        return "summarize"
    if any(w in q for w in ["fill in the blank", "fill in the blanks", "blanks", "fill the blank"]):
        return "fill_blanks"
    if any(w in q for w in ["quiz", "mcq", "multiple choice", "test me", "questions on",
                             "more question", "more questions", "create more", "generate more",
                             "another quiz", "different question", "new question",
                             "10 mcq", "15 mcq", "20 mcq", "5 mcq"]):
        return "quiz"
    if any(w in q for w in ["true", "false", "true/false", "true or false"]):
        return "true_false"
    if any(w in q for w in ["explain", "elaborate", "describe", "what is", "what are", "how does", "why"]):
        return "explain"
    return "general"


# Task-specific instruction blocks injected into the prompt
TASK_INSTRUCTIONS = {
    "summarize": """TASK: The user wants a SUMMARY. Follow this EXACT structure:

📌 OVERVIEW
- Write 2-3 simple lines summarizing the document
- Make it easy to understand for a beginner

---

🧠 KEY CONCEPTS
- Extract the most important concepts
- Format each as: **Keyword** — one simple explanation

---

🏭 HOW IT WORKS
- Break the process into clear steps
- Keep each step short and simple

---

🌍 REAL-WORLD APPLICATION
- Where is this used in real life?
- Keep it practical and relatable

---

⚠️ IMPORTANT INSIGHTS
- Highlight critical points or risks
- Use bullet points

---

💡 FINAL TAKEAWAYS
- List 3-5 key things the user must remember

---

🔥 KEY TERMS
- List important keywords separated by commas

STYLE RULES:
- Write detailed explanations, minimum 4-5 lines per section
- Use simple conversational English like a teacher explaining to a student
- Each bullet point must be a full sentence, not just keywords
- No long paragraphs
- Simple English like a human teacher
- Use **bold** for important words
- Use ONLY the document context,
- Do NOT stop early — cover all major points.""",

    "fill_blanks": """TASK: The user wants FILL-IN-THE-BLANK exercises.
- Generate exactly 3-5 fill-in-the-blank sentences.
- Use ONLY facts from the document context.
- Replace the key term or value with a blank: _______
- After all blanks, provide an "Answers:" section listing the correct words in order.
- Format:
  1. Sentence with _______.
  2. Sentence with _______.
  Answers: 1) word  
           2) word""",

    "quiz": """TASK: The user wants a QUIZ.
- The user may specify how many questions they want (e.g. "10 MCQs", "15 questions"). If specified, generate EXACTLY that many. If not specified, generate 5.
- Each question must have 4 options (A, B, C, D).
- Mark the correct answer clearly after each question.
- Do NOT stop early. Complete ALL requested questions.
- Format:
  Q1. Question text?
  A) ...
  B) ...
  C) ...
  D) ...
  Answer: X) ...""",

    "explain": """TASK: The user wants an EXPLANATION. Follow this EXACT structure:

🎯 HOOK
- Start with a simple question OR a relatable situation
- Make the user curious to read more

---

🧠 CORE EXPLANATION
- Explain the topic like teaching a beginner
- Use simple language
- Build understanding step by step
- Connect ideas naturally (no bullet points here — write in flowing sentences)

---

🔍 BREAK IT DOWN
- Explain each important part one by one
- For each part cover:
  • What it is
  • Why it matters
  • A simple example

---

⚙️ HOW IT WORKS
- Explain like a real-life process happening
- Use "Imagine this..." style explanation

---

💡 MAKE IT CLICK
- Give a real-world analogy
- Compare with something the user already knows

---

🚨 COMMON CONFUSION
- Mention 1-2 things people usually misunderstand
- Clarify them simply and clearly

---

✅ FINAL UNDERSTANDING
- End with a short 2-3 line summary of the whole idea
- Make it stick in the user's memory

STRICT RULES:
- Use ONLY the document context
- Simple English only — like a human teacher
- Never use robotic or academic language
- Do NOT skip any section""",

    "general": """TASK: Answer the user's question directly and concisely.
- Use ONLY the document context.
- Use bullet points if the answer has multiple parts.
- Keep the answer focused and free of padding.""",

    "notes": """TASK: Create DETAILED STUDY NOTES — minimum 15 lines.

Output format — use EXACTLY this structure:

📝 STUDY NOTES
━━━━━━━━━━━━━━

📌 TOPIC OVERVIEW
[2-3 sentences summarizing what this document is about]

📌 KEY CONCEPTS
- [Concept 1 — explain in 1-2 sentences]
- [Concept 2 — explain in 1-2 sentences]
- [Concept 3 — explain in 1-2 sentences]
- [Concept 4 — explain in 1-2 sentences]
- [Concept 5 — explain in 1-2 sentences]

📌 IMPORTANT FACTS & DETAILS
- [Fact 1 with specific detail from document]
- [Fact 2 with specific detail from document]
- [Fact 3 with specific detail from document]
- [Fact 4 with specific detail from document]

📌 DATES / NUMBERS / DATA (if any)
- [Any specific dates, numbers, statistics, or measurements from the document]
- [List all relevant data points found]

📌 KEY TERMS TO REMEMBER
- [Term 1]: [definition from document]
- [Term 2]: [definition from document]
- [Term 3]: [definition from document]

📌 CONCLUSIONS / OUTCOMES
- [What the document concludes or recommends]
- [Any final key point]

💡 KEY INSIGHT
[One sentence — the single most important thing to remember]

STRICT RULES:
- Minimum 15 bullet points across all sections
- Use ONLY information from the document context
- Do NOT write paragraphs — bullets ONLY
- Do NOT stop early""",

    "tutor": """TASK: You are an expert AI Tutor like Gemini.
- Explain the answer step by step with clear structure
- Use simple language with examples or analogies where helpful
- After answering, suggest 2 follow-up questions the student should explore
- End with a "💡 Key Insight:" one-line summary
- Be encouraging and educational in tone
- Still use ONLY the document context as your source""",
}

# ─── Background task ──────────────────────────────────────────────────────────
def process_document_background(
    file_id: str, file_location: str, filename: str, session_id: str
):
    global vector_db_store

    try:
        upload_status[file_id]["message"] = "Extracting text..."
        pages = extract_text_with_pages(file_location, filename)

        if not pages:
            upload_status[file_id] = {
                "status": "error",
                "message": "No readable text found in file.",
                "chunks": 0,
            }
            return

        upload_status[file_id]["message"] = "Splitting into chunks..."
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=300,
            chunk_overlap=30,
        )
        all_texts: List[str] = []
        all_metadatas: List[dict] = []

        for page_data in pages:
            page_chunks = splitter.split_text(page_data["text"])
            for chunk in page_chunks:
                all_texts.append(chunk)
                all_metadatas.append({"page": page_data["page"]})

        upload_status[file_id]["message"] = "Building FAISS index..."
        BATCH_SIZE = 64
        session_vdb = None

        for i in range(0, len(all_texts), BATCH_SIZE):
            batch_texts = all_texts[i:i + BATCH_SIZE]
            batch_metas = all_metadatas[i:i + BATCH_SIZE]
            batch_vecs  = embeddings.embed_documents(batch_texts)

            if session_vdb is None:
                session_vdb = FAISS.from_embeddings(
                    list(zip(batch_texts, batch_vecs)),
                    embeddings,
                    metadatas=batch_metas,
                )
            else:
                session_vdb.add_embeddings(
                    list(zip(batch_texts, batch_vecs)),
                    metadatas=batch_metas,
                )

            pct = min(99, int((i + BATCH_SIZE) / len(all_texts) * 100))
            upload_status[file_id]["message"] = f"Indexing… {pct}%"

        # ← Store under this user's session
        vector_db_store[session_id] = session_vdb

        upload_status[file_id] = {
            "status": "ready",
            "message": f"{filename} processed successfully.",
            "chunks": len(all_texts),
        }

    except Exception as e:
        upload_status[file_id] = {
            "status": "error",
            "message": str(e),
            "chunks": 0,
        }


# ─── Routes ───────────────────────────────────────────────────────────────────
@app.get("/")
def home():
    return {"message": "DocuMind AI backend is running 🚀"}


@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    session_id: Optional[str] = Form(None),
):
    supported = (".pdf", ".jpg", ".jpeg", ".png", ".docx")
    if not any(file.filename.endswith(ext) for ext in supported):
        return {"error": "Unsupported file format"}

    file_id = str(uuid.uuid4())
    sid = session_id or str(uuid.uuid4())
    file_location = f"uploaded_{file.filename}"

    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    upload_status[file_id] = {
        "status": "processing",
        "message": "File received. Processing in background...",
        "chunks": 0,
    }

    background_tasks.add_task(
        process_document_background, file_id, file_location, file.filename, sid
    )

    return {
        "status": "processing",
        "file_id": file_id,
        "session_id": sid,
        "message": "File uploaded. Processing in background.",
        "filename": file.filename,
    }


@app.get("/status/{file_id}")
def get_status(file_id: str):
    if file_id not in upload_status:
        return {"error": "Unknown file_id"}
    return upload_status[file_id]


@app.post("/ask")
def ask_question(body: AskRequest):
    query = body.query.strip()
    sid = body.session_id or ""
    vector_db = vector_db_store.get(sid)

    if vector_db is None:
        return {
            "answer": "Please upload a document first 🙂",
            "general": True,
            "confidence": "low",
            "sources": [],
        }

    docs = vector_db.similarity_search(query, k=8)
    context = "\n\n".join([doc.page_content for doc in docs])

    sources = [
        {
            "page": doc.metadata.get("page", 1),
            "text": doc.page_content[:200].strip(),
        }
        for doc in docs
    ]

    if len(context.strip()) < 50:
        fallback_answer = call_ollama(
            f"You are a helpful assistant. Answer in 1-2 lines. "
            f"Do NOT make up facts.\n\nQuestion: {query}\nAnswer:"
        )
        return {
            "answer": f"Not found in document.\n\nGeneral knowledge: {fallback_answer}",
            "general": True,
            "confidence": "low",
            "sources": [],
        }

    history_block = format_history(body.history or [], max_turns=5)
    history_section = (
        f"### Conversation History\n{history_block}\n\n"
        if history_block
        else ""
    )

    task = detect_task(query)
    if body.tutor_mode:
        task_instruction = TASK_INSTRUCTIONS["tutor"]
    else:
        task_instruction = TASK_INSTRUCTIONS[task]

    main_prompt = f"""You are DocuMind AI, a strict document-based assistant.
Your sole source of truth is the Document Context below.

=== CORE RULES ===
1. You MUST follow the task instruction exactly (see TASK section below).
2. You MUST base your answer ONLY on the Document Context provided.
3. Do NOT hallucinate or add facts not present in the context.
4. If the answer cannot be found in the context, say EXACTLY:
   "Not found in document."
   Then give a short general answer (max 2 lines) labelled "General knowledge:".
5. Do NOT ignore task instructions under any circumstances.
6. Format output clearly — use bullet points or numbered lists where appropriate.

=== TASK ===
{task_instruction}

{history_section}=== Document Context ===
{context}

=== User Question ===
{query}

=== Answer ===
"""

    answer = call_ollama(main_prompt)

    if (answer.upper().startswith("NOT FOUND")
        or "NOT_FOUND" in answer.upper()
        or "not found in document" in answer.lower()):
        fallback_answer = call_ollama(
            f"You are a helpful assistant. Answer in 1-2 lines. "
            f"Do NOT make up facts.\n\nQuestion: {query}\nAnswer:"
        )
        return {
            "answer": f"Not found in document.\n\nGeneral knowledge: {fallback_answer}",
            "general": True,
            "confidence": "low",
            "sources": [],
        }

    confidence = score_confidence(context, answer)
    return {
        "answer": answer,
        "general": False,
        "confidence": confidence,
        "sources": sources,
    }